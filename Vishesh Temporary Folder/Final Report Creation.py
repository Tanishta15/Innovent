import pandas as pd
from datetime import datetime, timedelta
import openai
from docx import Document
from docx.shared import Inches

# Load and filter CSV data
data = pd.read_csv('/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/results.csv')
data['Date'] = pd.to_datetime(data['Date'])
Product_ID = data['Product_ID'].unique()

part_to_map_image = {
    'A1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/MapA1.png',
    'B1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/MapB1.png',
    'C1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/MapC1.png'
}

part_to_direction = {
    'A1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/directionsA1.txt',
    'B1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/directionsB1.txt',
    'C1': '/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/directionsC1.txt'
}

def generate_report_summary(filtered_data):
    openai.api_key = 'sk-None-DiN49NbLRnkIQwVgTdyBT3BlbkFJPuIv6N7cbCTw1RpyDwGs'
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # Or the model you are using
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": f"Generate a summary report from the following inventory data:\n{filtered_data}"}
        ],
        max_tokens=500
    )
    
    return response['choices'][0]['message']['content'].strip()

def read_directions_file(file_path):
    with open(file_path, 'r') as file:
        directions = file.read()
    return directions

def filter_non_latin_characters(text):
    return text.encode('latin-1', 'ignore').decode('latin-1')

def create_docx_report(summary, part_to_map_image, relevant_parts, part_to_direction):
    doc = Document()

    # Title
    doc.add_heading('Inventory Management Report', level=1)

    # Summary
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)

    # Add Relevant Map Images and Instructions
    for part in relevant_parts:
        if part in part_to_map_image:
            doc.add_heading(f'Map for Part {part}', level=3)
            doc.add_picture(part_to_map_image[part], width=Inches(5))

        if part in part_to_direction:
            doc.add_heading(f'Directions for Part {part}', level=3)
            instructions = read_directions_file(part_to_direction[part])
            doc.add_paragraph(instructions)

    # Save the document
    doc.save('/Users/visheshgoyal/Innovent/Vishesh Temporary Folder/inventory_management_report.docx')

summary = generate_report_summary(data)
create_docx_report(summary, part_to_map_image, Product_ID, part_to_direction)